"""
Script mejorado para insertar referencias a capturas de pantalla en el documento ManualUsuario.docx.
Este script busca los marcadores [CAPTURA DE PANTALLA X] en el documento y los reemplaza con
referencias formateadas y espacios para las imágenes.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import re

def insertar_capturas_en_docx():
    doc_path = "ManualUsuario.docx"
    
    if not os.path.exists(doc_path):
        print(f"Error: No se encontró el archivo {doc_path}")
        return
    
    print(f"Abriendo documento: {doc_path}")
    doc = Document(doc_path)
    
    # Mapeo de números de captura a archivos y descripciones
    capturas_info = {
        1: {
            "archivo": "capturas-manual/captura-01-panel-general.jpg",
            "descripcion": "Pantalla principal del sistema mostrando el menú de navegación y el panel general con gráficos interactivos"
        },
        2: {
            "archivo": "capturas-manual/captura-02-boton-autenticacion.jpg",
            "descripcion": "Vista del header mostrando el botón de autenticación '👤 Iniciar Sesión / Registrarse' en la esquina superior derecha"
        },
        3: {
            "archivo": "capturas-manual/captura-03-modal-registro.jpg",
            "descripcion": "Modal de autenticación mostrando las pestañas de 'Iniciar Sesión' y 'Registrarse', con la pestaña 'Registrarse' activa"
        },
        4: {
            "archivo": "capturas-manual/captura-04-formulario-registro.jpg",
            "descripcion": "Formulario de registro completo con los campos de Nombre de Usuario, Email y Contraseña, mostrando validación en tiempo real"
        },
        5: {
            "archivo": "capturas-manual/captura-05-header-autenticado.jpg",
            "descripcion": "Header después de iniciar sesión mostrando el nombre del usuario y el botón 'Salir'"
        },
        6: {
            "archivo": "capturas-manual/captura-06-modal-login.jpg",
            "descripcion": "Modal de autenticación con la pestaña 'Iniciar Sesión' activa, mostrando los campos de nombre de usuario, email y contraseña"
        },
        7: {
            "archivo": "capturas-manual/captura-07-menu-principal.jpg",
            "descripcion": "Vista completa del menú principal de navegación mostrando todos los botones: Panel General, Búsqueda Avanzada, Modo Fácil, Comparar, Mi Cuenta"
        },
        8: {
            "archivo": "capturas-manual/captura-08-panel-graficos.jpg",
            "descripcion": "Vista del Panel General mostrando los gráficos interactivos de Precio Promedio por Marca (MXN) y Capacidad de Batería (mAh)"
        },
        9: {
            "archivo": "capturas-manual/captura-09-busqueda-filtros.jpg",
            "descripcion": "Vista de la sección de Búsqueda Avanzada mostrando todos los filtros disponibles organizados en secciones"
        },
        10: {
            "archivo": "capturas-manual/captura-10-resultados-busqueda.jpg",
            "descripcion": "Resultados de búsqueda mostrando tarjetas de smartphones con imágenes, especificaciones, precios y botones de acción"
        },
        11: {
            "archivo": "capturas-manual/captura-11-modo-facil-presupuesto.jpg",
            "descripcion": "Vista del Modo Fácil mostrando la pregunta '¿Cuál es tu presupuesto aproximado?' con opciones de presupuesto"
        },
        12: {
            "archivo": "capturas-manual/captura-12-modo-facil-uso.jpg",
            "descripcion": "Vista del Modo Fácil mostrando la pregunta '¿Para qué usarás principalmente tu teléfono?' con opciones de uso"
        },
        13: {
            "archivo": "capturas-manual/captura-13-resultados-modo-facil.jpg",
            "descripcion": "Resultados del Modo Fácil mostrando '¡Recomendaciones Perfectas!' con smartphones rankeados y porcentajes de coincidencia"
        },
        14: {
            "archivo": "capturas-manual/captura-14-tabla-comparativa.jpg",
            "descripcion": "Vista de la sección Comparar mostrando mensaje 'No hay teléfonos para comparar' con instrucciones"
        },
        15: {
            "archivo": "capturas-manual/captura-15-boton-flotante.jpg",
            "descripcion": "Botón flotante de comparación en la esquina inferior derecha"
        },
        16: {
            "archivo": "capturas-manual/captura-16-mi-cuenta-vista.jpg",
            "descripcion": "Vista de la sección Mi Cuenta mostrando el header con estadísticas y las pestañas de Favoritos e Historial"
        },
        17: {
            "archivo": "capturas-manual/captura-17-favoritos-vacio.jpg",
            "descripcion": "Pestaña de Favoritos mostrando estado vacío con mensaje y botón 'Explorar Smartphones'"
        },
        18: {
            "archivo": "capturas-manual/captura-18-historial.jpg",
            "descripcion": "Pestaña de Historial de búsquedas con búsquedas guardadas"
        },
        19: {
            "archivo": "capturas-manual/captura-19-busqueda-tiempo-real.jpg",
            "descripcion": "Barra de búsqueda 'Buscar por nombre' con resultados apareciendo en tiempo real"
        },
        20: {
            "archivo": "capturas-manual/captura-20-filtros-aplicados.jpg",
            "descripcion": "Vista de los filtros aplicados en Búsqueda Avanzada con resultados correspondientes"
        },
        21: {
            "archivo": "capturas-manual/captura-21-boton-favoritos.jpg",
            "descripcion": "Tarjeta de smartphone mostrando el botón de favoritos (❤️) y comparar (📊) en la esquina superior derecha"
        },
        22: {
            "archivo": "capturas-manual/captura-22-boton-comparar.jpg",
            "descripcion": "Botón de comparar (📊) en tarjeta y botón flotante con contador"
        },
        23: {
            "archivo": "capturas-manual/captura-23-modal-comentarios.jpg",
            "descripcion": "Modal de detalles de smartphone mostrando sección '💬 Comentarios y Reseñas' con formulario"
        },
        24: {
            "archivo": "capturas-manual/captura-24-lista-comentarios.jpg",
            "descripcion": "Sección de comentarios mostrando lista de reseñas con calificaciones, nombres de usuarios y fechas"
        },
        25: {
            "archivo": "capturas-manual/captura-25-boton-chatbot.jpg",
            "descripcion": "Botón flotante del chatbot (💬) en la esquina inferior izquierda"
        },
        26: {
            "archivo": "capturas-manual/captura-26-ventana-chatbot.jpg",
            "descripcion": "Ventana del chatbot abierta con mensaje de bienvenida y opciones rápidas"
        }
    }
    
    # Buscar y reemplazar marcadores en todos los párrafos
    cambios_realizados = 0
    
    for para in doc.paragraphs:
        texto = para.text
        
        # Buscar patrones como [CAPTURA DE PANTALLA 1: ...] o [CAPTURA DE PANTALLA 1]
        patron = r'\[CAPTURA DE PANTALLA (\d+)(?::.*?)?\]'
        matches = re.finditer(patron, texto, re.IGNORECASE)
        
        for match in list(matches):
            num_captura = int(match.group(1))
            
            if num_captura in capturas_info:
                info = capturas_info[num_captura]
                
                # Limpiar el párrafo actual
                para.clear()
                
                # Agregar título de la figura
                run_titulo = para.add_run(f"Figura {num_captura}: {info['descripcion']}")
                run_titulo.bold = True
                run_titulo.font.size = Pt(11)
                
                # Agregar salto de línea
                para.add_run("\n")
                
                # Intentar insertar la imagen si existe
                if os.path.exists(info['archivo']):
                    try:
                        para.add_run().add_picture(info['archivo'], width=Inches(5.5))
                        print(f"✓ Imagen {num_captura} insertada: {info['archivo']}")
                    except Exception as e:
                        run_img = para.add_run(f"[IMAGEN: {info['archivo']}]")
                        run_img.italic = True
                        run_img.font.color.rgb = RGBColor(128, 128, 128)
                        print(f"⚠ Error al insertar imagen {num_captura}: {e}")
                else:
                    # Agregar marcador de imagen faltante
                    run_img = para.add_run(f"[IMAGEN: {info['archivo']}]")
                    run_img.italic = True
                    run_img.font.color.rgb = RGBColor(255, 0, 0)
                    print(f"⚠ Imagen {num_captura} no encontrada: {info['archivo']}")
                
                # Centrar el párrafo
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                cambios_realizados += 1
    
    # Guardar documento
    doc.save(doc_path)
    print(f"\n✓ Documento actualizado: {doc_path}")
    print(f"Total de cambios realizados: {cambios_realizados}")
    
    if cambios_realizados == 0:
        print("\n⚠ No se encontraron marcadores [CAPTURA DE PANTALLA X] en el documento.")
        print("El documento puede no tener los marcadores o ya estar actualizado.")

if __name__ == "__main__":
    try:
        insertar_capturas_en_docx()
    except ImportError:
        print("Error: Se requiere la librería python-docx")
        print("Instálala con: pip install python-docx")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

