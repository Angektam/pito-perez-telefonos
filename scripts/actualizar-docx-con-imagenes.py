"""
Script para actualizar el documento ManualUsuario.docx con referencias a las capturas de pantalla.
Este script agrega texto descriptivo donde deben ir las imágenes.
Para insertar las imágenes reales, ábrelo en Word y reemplaza los textos con las imágenes.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def actualizar_documento():
    # Ruta del documento
    doc_path = "ManualUsuario.docx"
    
    # Verificar si existe
    if not os.path.exists(doc_path):
        print(f"Error: No se encontró el archivo {doc_path}")
        print("Creando un nuevo documento...")
        doc = Document()
    else:
        print(f"Abriendo documento existente: {doc_path}")
        doc = Document(doc_path)
    
    # Lista de capturas con sus descripciones
    capturas = [
        {
            "num": 1,
            "titulo": "Pantalla principal del sistema",
            "descripcion": "Vista principal del sistema Pito Pérez mostrando el menú de navegación y el panel general con gráficos interactivos de Precio por Marca y Capacidad de Batería",
            "archivo": "capturas-manual/captura-01-panel-general.jpg"
        },
        {
            "num": 2,
            "titulo": "Botón de autenticación",
            "descripcion": "Vista del header mostrando el botón '👤 Iniciar Sesión / Registrarse' en la esquina superior derecha",
            "archivo": "capturas-manual/captura-02-boton-autenticacion.jpg"
        },
        {
            "num": 3,
            "titulo": "Modal de registro",
            "descripcion": "Modal de autenticación mostrando las pestañas de 'Iniciar Sesión' y 'Registrarse', con la pestaña 'Registrarse' activa",
            "archivo": "capturas-manual/captura-03-modal-registro.jpg"
        },
        {
            "num": 4,
            "titulo": "Formulario de registro",
            "descripcion": "Formulario de registro completo con los campos de Nombre de Usuario, Email y Contraseña",
            "archivo": "capturas-manual/captura-04-formulario-registro.jpg"
        },
        {
            "num": 5,
            "titulo": "Header autenticado",
            "descripcion": "Header después de iniciar sesión mostrando el nombre del usuario y el botón 'Salir'",
            "archivo": "capturas-manual/captura-05-header-autenticado.jpg"
        },
        {
            "num": 6,
            "titulo": "Modal de inicio de sesión",
            "descripcion": "Modal de autenticación con la pestaña 'Iniciar Sesión' activa, mostrando los campos de nombre de usuario, email y contraseña",
            "archivo": "capturas-manual/captura-06-modal-login.jpg"
        },
        {
            "num": 7,
            "titulo": "Menú principal",
            "descripcion": "Vista completa del menú principal de navegación mostrando todos los botones: Panel General, Búsqueda Avanzada, Modo Fácil, Comparar, Mi Cuenta",
            "archivo": "capturas-manual/captura-07-menu-principal.jpg"
        },
        {
            "num": 8,
            "titulo": "Panel General con gráficos",
            "descripcion": "Vista del Panel General mostrando gráficos interactivos de Precio Promedio por Marca (MXN) y Capacidad de Batería (mAh)",
            "archivo": "capturas-manual/captura-08-panel-graficos.jpg"
        },
        {
            "num": 9,
            "titulo": "Búsqueda Avanzada con filtros",
            "descripcion": "Vista de la sección de Búsqueda Avanzada mostrando todos los filtros disponibles organizados en secciones: Información Básica, Especificaciones, y Rendimiento y Precio",
            "archivo": "capturas-manual/captura-09-busqueda-filtros.jpg"
        },
        {
            "num": 10,
            "titulo": "Resultados de búsqueda",
            "descripcion": "Resultados de búsqueda mostrando tarjetas de smartphones (Galaxy A34, A54, A55) con imágenes, especificaciones, precios y botones de acción",
            "archivo": "capturas-manual/captura-10-resultados-busqueda.jpg"
        },
        {
            "num": 11,
            "titulo": "Modo Fácil - Presupuesto",
            "descripcion": "Vista del Modo Fácil mostrando la pregunta '¿Cuál es tu presupuesto aproximado?' con opciones: Económico, Accesible, Intermedio, Alto, Premium, Flexible",
            "archivo": "capturas-manual/captura-11-modo-facil-presupuesto.jpg"
        },
        {
            "num": 12,
            "titulo": "Modo Fácil - Tipo de uso",
            "descripcion": "Vista del Modo Fácil mostrando la pregunta '¿Para qué usarás principalmente tu teléfono?' con opciones: Básico, Redes Sociales, Juegos, Trabajo, Creativo, Estudios, Viajes, Uso Mixto",
            "archivo": "capturas-manual/captura-12-modo-facil-uso.jpg"
        },
        {
            "num": 13,
            "titulo": "Resultados del Modo Fácil",
            "descripcion": "Resultados del Modo Fácil mostrando '¡Recomendaciones Perfectas!' con smartphones rankeados (iPhone 16 Pro Max, iPhone 16 Pro, iPhone 16) y porcentajes de coincidencia (80%)",
            "archivo": "capturas-manual/captura-13-resultados-modo-facil.jpg"
        },
        {
            "num": 14,
            "titulo": "Vista de comparación",
            "descripcion": "Vista de la sección Comparar mostrando mensaje 'No hay teléfonos para comparar' con instrucciones y botones para ir a Búsqueda Avanzada o Modo Fácil",
            "archivo": "capturas-manual/captura-14-tabla-comparativa.jpg"
        },
        {
            "num": 15,
            "titulo": "Botón flotante de comparación",
            "descripcion": "Botón flotante de comparación en la esquina inferior derecha (se muestra cuando hay smartphones agregados a la comparación)",
            "archivo": "capturas-manual/captura-15-boton-flotante.jpg"
        },
        {
            "num": 16,
            "titulo": "Mi Cuenta - Vista general",
            "descripcion": "Vista de la sección Mi Cuenta mostrando el header con estadísticas (0 Favoritos, 0 Búsquedas) y las pestañas de Favoritos e Historial",
            "archivo": "capturas-manual/captura-16-mi-cuenta-vista.jpg"
        },
        {
            "num": 17,
            "titulo": "Pestaña de Favoritos",
            "descripcion": "Pestaña de Favoritos mostrando estado vacío con mensaje 'No tienes favoritos aún' y botón 'Explorar Smartphones'",
            "archivo": "capturas-manual/captura-17-favoritos-vacio.jpg"
        },
        {
            "num": 18,
            "titulo": "Pestaña de Historial",
            "descripcion": "Pestaña de Historial de búsquedas (se muestra cuando hay búsquedas guardadas con filtros aplicados)",
            "archivo": "capturas-manual/captura-18-historial.jpg"
        },
        {
            "num": 19,
            "titulo": "Búsqueda en tiempo real",
            "descripcion": "Barra de búsqueda 'Buscar por nombre' con resultados apareciendo en tiempo real mientras se escribe",
            "archivo": "capturas-manual/captura-19-busqueda-tiempo-real.jpg"
        },
        {
            "num": 20,
            "titulo": "Filtros aplicados",
            "descripcion": "Vista de los filtros aplicados en Búsqueda Avanzada con resultados correspondientes actualizándose",
            "archivo": "capturas-manual/captura-20-filtros-aplicados.jpg"
        },
        {
            "num": 21,
            "titulo": "Botón de favoritos",
            "descripcion": "Tarjeta de smartphone mostrando el botón de favoritos (❤️) y comparar (📊) en la esquina superior derecha",
            "archivo": "capturas-manual/captura-21-boton-favoritos.jpg"
        },
        {
            "num": 22,
            "titulo": "Botón de comparar",
            "descripcion": "Botón de comparar (📊) en tarjeta y botón flotante con contador (se muestra cuando hay smartphones agregados)",
            "archivo": "capturas-manual/captura-22-boton-comparar.jpg"
        },
        {
            "num": 23,
            "titulo": "Modal con comentarios",
            "descripcion": "Modal de detalles de smartphone (Galaxy A34) mostrando sección '💬 Comentarios y Reseñas' con calificación promedio (0.0) y mensaje para iniciar sesión",
            "archivo": "capturas-manual/captura-23-modal-comentarios.jpg"
        },
        {
            "num": 24,
            "titulo": "Lista de comentarios",
            "descripcion": "Sección de comentarios mostrando mensaje 'Aún no hay comentarios sobre este teléfono. ¡Sé el primero en compartir tu opinión!'",
            "archivo": "capturas-manual/captura-24-lista-comentarios.jpg"
        },
        {
            "num": 25,
            "titulo": "Botón del chatbot",
            "descripcion": "Botón flotante del chatbot (💬) en la esquina inferior izquierda",
            "archivo": "capturas-manual/captura-25-boton-chatbot.jpg"
        },
        {
            "num": 26,
            "titulo": "Ventana del chatbot",
            "descripcion": "Ventana del chatbot abierta con mensaje de bienvenida y opciones rápidas",
            "archivo": "capturas-manual/captura-26-ventana-chatbot.jpg"
        }
    ]
    
    # Agregar título al documento si es nuevo
    if len(doc.paragraphs) == 0:
        title = doc.add_heading('Manual de Usuario del Sistema Pito Pérez', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Agregar sección de capturas de pantalla
    doc.add_heading('Índice de Capturas de Pantalla', 1)
    
    doc.add_paragraph(
        'Este documento incluye referencias a 26 capturas de pantalla que ilustran '
        'las funcionalidades del sistema. Para insertar las imágenes reales, reemplaza '
        'cada sección marcada con [CAPTURA X] con la imagen correspondiente desde la '
        'carpeta capturas-manual/.'
    )
    
    # Agregar cada captura
    for captura in capturas:
        # Título de la captura
        heading = doc.add_heading(f'Captura de Pantalla {captura["num"]}: {captura["titulo"]}', 2)
        
        # Descripción
        p = doc.add_paragraph(captura["descripcion"])
        p.style = 'List Bullet'
        
        # Marcador para la imagen
        p_img = doc.add_paragraph()
        p_img.add_run(f'[CAPTURA {captura["num"]}]').bold = True
        p_img.add_run(f'\nArchivo: {captura["archivo"]}\n')
        p_img.add_run('→ Inserta aquí la imagen correspondiente')
        p_img.italic = True
        
        # Intentar insertar la imagen si existe
        img_path = captura["archivo"]
        if os.path.exists(img_path):
            try:
                run = p_img.runs[0]
                # Agregar imagen
                doc.add_picture(img_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"✓ Imagen {captura['num']} insertada: {img_path}")
            except Exception as e:
                print(f"⚠ No se pudo insertar imagen {captura['num']}: {e}")
        else:
            print(f"⚠ Imagen no encontrada: {img_path}")
        
        # Espaciado
        doc.add_paragraph()
    
    # Guardar documento
    output_path = "ManualUsuario.docx"
    doc.save(output_path)
    print(f"\n✓ Documento guardado: {output_path}")
    print(f"\nTotal de capturas procesadas: {len(capturas)}")
    print("\nNota: Si alguna imagen no se insertó, ábre el documento en Word y")
    print("reemplaza los marcadores [CAPTURA X] con las imágenes correspondientes.")

if __name__ == "__main__":
    try:
        actualizar_documento()
    except ImportError:
        print("Error: Se requiere la librería python-docx")
        print("Instálala con: pip install python-docx")
    except Exception as e:
        print(f"Error: {e}")

