"""
Script para generar un documento .docx completo del manual de usuario
con las referencias a las capturas de pantalla insertadas en los lugares correctos.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generar_docx_completo():
    # Leer el contenido del ManualUsuario.txt
    txt_path = "ManualUsuario.txt"
    if not os.path.exists(txt_path):
        print(f"Error: No se encontró {txt_path}")
        return
    
    print(f"Leyendo contenido de {txt_path}...")
    with open(txt_path, 'r', encoding='utf-8') as f:
        contenido = f.read()
    
    # Crear nuevo documento
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Mapeo de capturas
    capturas_map = {
        1: "capturas-manual/captura-01-panel-general.jpg",
        2: "capturas-manual/captura-02-boton-autenticacion.jpg",
        3: "capturas-manual/captura-03-modal-registro.jpg",
        4: "capturas-manual/captura-04-formulario-registro.jpg",
        5: "capturas-manual/captura-05-header-autenticado.jpg",
        6: "capturas-manual/captura-06-modal-login.jpg",
        7: "capturas-manual/captura-07-menu-principal.jpg",
        8: "capturas-manual/captura-08-panel-graficos.jpg",
        9: "capturas-manual/captura-09-busqueda-filtros.jpg",
        10: "capturas-manual/captura-10-resultados-busqueda.jpg",
        11: "capturas-manual/captura-11-modo-facil-presupuesto.jpg",
        12: "capturas-manual/captura-12-modo-facil-uso.jpg",
        13: "capturas-manual/captura-13-resultados-modo-facil.jpg",
        14: "capturas-manual/captura-14-tabla-comparativa.jpg",
        15: "capturas-manual/captura-15-boton-flotante.jpg",
        16: "capturas-manual/captura-16-mi-cuenta-vista.jpg",
        17: "capturas-manual/captura-17-favoritos-vacio.jpg",
        18: "capturas-manual/captura-18-historial.jpg",
        19: "capturas-manual/captura-19-busqueda-tiempo-real.jpg",
        20: "capturas-manual/captura-20-filtros-aplicados.jpg",
        21: "capturas-manual/captura-21-boton-favoritos.jpg",
        22: "capturas-manual/captura-22-boton-comparar.jpg",
        23: "capturas-manual/captura-23-modal-comentarios.jpg",
        24: "capturas-manual/captura-24-lista-comentarios.jpg",
        25: "capturas-manual/captura-25-boton-chatbot.jpg",
        26: "capturas-manual/captura-26-ventana-chatbot.jpg"
    }
    
    descripciones = {
        1: "Pantalla principal del sistema mostrando el menú de navegación y el panel general con gráficos interactivos",
        2: "Vista del header mostrando el botón de autenticación '👤 Iniciar Sesión / Registrarse'",
        3: "Modal de autenticación con pestañas 'Iniciar Sesión' y 'Registrarse'",
        4: "Formulario de registro completo con campos de Nombre de Usuario, Email y Contraseña",
        5: "Header después de iniciar sesión mostrando el nombre del usuario",
        6: "Modal de autenticación con pestaña 'Iniciar Sesión' activa",
        7: "Vista completa del menú principal de navegación",
        8: "Panel General mostrando gráficos interactivos de Precio por Marca y Batería",
        9: "Sección de Búsqueda Avanzada mostrando todos los filtros disponibles",
        10: "Resultados de búsqueda mostrando tarjetas de smartphones",
        11: "Modo Fácil - Pregunta sobre presupuesto aproximado",
        12: "Modo Fácil - Pregunta sobre tipo de uso del teléfono",
        13: "Resultados del Modo Fácil con recomendaciones personalizadas",
        14: "Vista de la sección Comparar con instrucciones",
        15: "Botón flotante de comparación en la esquina inferior derecha",
        16: "Sección Mi Cuenta con pestañas de Favoritos e Historial",
        17: "Pestaña de Favoritos mostrando estado vacío",
        18: "Pestaña de Historial de búsquedas",
        19: "Barra de búsqueda con resultados en tiempo real",
        20: "Filtros aplicados en Búsqueda Avanzada con resultados",
        21: "Tarjeta de smartphone con botones de favoritos y comparar",
        22: "Botón de comparar y botón flotante con contador",
        23: "Modal de detalles de smartphone con sección de comentarios",
        24: "Lista de comentarios y reseñas de usuarios",
        25: "Botón flotante del chatbot en la esquina inferior izquierda",
        26: "Ventana del chatbot abierta con mensaje de bienvenida"
    }
    
    # Dividir contenido en líneas
    lineas = contenido.split('\n')
    
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        
        # Buscar marcadores de captura
        import re
        match = re.search(r'\[CAPTURA DE PANTALLA (\d+)(?::.*?)?\]', linea, re.IGNORECASE)
        
        if match:
            num_captura = int(match.group(1))
            
            # Agregar título de la figura
            para = doc.add_paragraph()
            run = para.add_run(f"Figura {num_captura}: {descripciones.get(num_captura, '')}")
            run.bold = True
            run.font.size = Pt(11)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Intentar insertar imagen
            img_path = capturas_map.get(num_captura)
            if img_path and os.path.exists(img_path):
                try:
                    para_img = doc.add_paragraph()
                    para_img.add_run().add_picture(img_path, width=Inches(5.5))
                    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"✓ Imagen {num_captura} insertada")
                except Exception as e:
                    para_img = doc.add_paragraph()
                    run_img = para_img.add_run(f"[IMAGEN: {img_path}]")
                    run_img.italic = True
                    run_img.font.color.rgb = RGBColor(128, 128, 128)
                    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"⚠ Error al insertar imagen {num_captura}: {e}")
            else:
                para_img = doc.add_paragraph()
                run_img = para_img.add_run(f"[IMAGEN: {img_path or 'No especificada'}]")
                run_img.italic = True
                run_img.font.color.rgb = RGBColor(255, 0, 0)
                para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"⚠ Imagen {num_captura} no encontrada: {img_path}")
            
            # Agregar espacio
            doc.add_paragraph()
        else:
            # Agregar línea normal
            if linea.strip():
                # Detectar títulos
                if linea.startswith('='):
                    continue  # Saltar separadores
                elif re.match(r'^\d+\.', linea):  # Título de sección
                    doc.add_heading(linea.strip(), level=1)
                elif re.match(r'^\d+\.\d+', linea):  # Subtítulo
                    doc.add_heading(linea.strip(), level=2)
                elif re.match(r'^\d+\.\d+\.\d+', linea):  # Sub-subtítulo
                    doc.add_heading(linea.strip(), level=3)
                else:
                    doc.add_paragraph(linea.strip())
        
        i += 1
    
    # Guardar documento
    output_path = "ManualUsuario.docx"
    doc.save(output_path)
    print(f"\n✓ Documento generado: {output_path}")
    print(f"Total de capturas procesadas: {len(capturas_map)}")

if __name__ == "__main__":
    try:
        generar_docx_completo()
    except ImportError:
        print("Error: Se requiere la librería python-docx")
        print("Instálala con: pip install python-docx")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

